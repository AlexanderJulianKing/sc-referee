#!/usr/bin/env python3
"""Build consolidated review copies and integrity metadata for the specification."""
from __future__ import annotations

import hashlib
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
VERSION = "0.5.0-draft"
SCHEMA_VERSION = "0.5.0"
DATE = "2026-07-27"
DOCX_NAME = "sc-referee-specification-v0.5.0.docx"

DOCS = sorted((ROOT / "docs").glob("[0-9][0-9]-*.md"))
ADRS = sorted((ROOT / "adrs").glob("ADR-*.md"))


def strip_document_title(text: str, expected: str | None = None) -> str:
    """Return normalized Markdown while preserving the document's first heading."""
    value = text.strip()
    if expected and not value.startswith(expected):
        raise ValueError(f"Expected {expected!r} at document start")
    return value + "\n"


def rebase_relative_links(text: str, prefix: str) -> str:
    return re.sub(
        r"\]\((?!https?://|mailto:|#)([^)]+)\)",
        lambda m: "](" + prefix + m.group(1) + ")",
        text,
    )


def _heading_with_id(text: str, anchor: str) -> str:
    lines = text.strip().splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Expected top-level Markdown heading")
    lines[0] = f"{lines[0]} {{#{anchor}}}"
    return "\n".join(lines).strip() + "\n"


def _demote_first_heading(text: str) -> str:
    lines = text.strip().splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Expected top-level Markdown heading")
    lines[0] = "#" + lines[0]
    return "\n".join(lines).strip() + "\n"


def build_master() -> Path:
    module_entries: list[tuple[str, str, str]] = []
    for path in DOCS:
        raw = path.read_text(encoding="utf-8")
        title = raw.splitlines()[0][2:].strip()
        chapter_num = path.name.split("-", 1)[0]
        module_entries.append((title, f"chapter-{chapter_num}", raw))

    contents = [
        ("Accepted policy decisions for version 0.5", "accepted-decisions"),
        *[(title, anchor) for title, anchor, _ in module_entries],
        ("Appendix A. Acceptance criteria", "appendix-a"),
        ("Appendix B. Architecture Decision Records", "appendix-b"),
        ("Appendix C. Reference index", "appendix-c"),
    ]

    parts = [
        "---\n",
        "title: sc-referee Scientific Audit Specification\n",
        f"subtitle: Version {VERSION} - working draft\n",
        "author: sc-referee design project\n",
        f"date: {DATE}\n",
        "lang: en-US\n",
        "---\n\n",
        "# Contents {#contents}\n\n",
    ]
    for title, anchor in contents:
        parts.append(f"- [{title}](#{anchor})\n")
    parts.extend([
        "\n> **Review status:** This is a design specification, not a claim that the system has been implemented or scientifically validated. Open decisions remain normative gaps until accepted through an ADR.\n\n",
        "The modular Markdown documents in `docs/`, the ADRs, the accepted-decision log, and the machine registers are the editing sources of truth. This consolidated file is generated for review.\n\n",
        "# Accepted policy decisions for version 0.5 {#accepted-decisions}\n\n",
    ])

    decisions = (ROOT / "DECISIONS_v0.5.md").read_text(encoding="utf-8")
    decisions = decisions.replace("# Accepted decisions in specification 0.5.0", "", 1).strip()
    parts.append(decisions + "\n")

    for _title, anchor, raw in module_entries:
        parts.append("\n\n")
        parts.append(_heading_with_id(raw, anchor))

    parts.append("\n\n# Appendix A. Acceptance criteria {#appendix-a}\n\n")
    ac_text = (ROOT / "ACCEPTANCE_CRITERIA.md").read_text(encoding="utf-8")
    ac_text = ac_text.replace("# Initial acceptance criteria", "", 1).strip()
    parts.append(ac_text + "\n")

    parts.append("\n\n# Appendix B. Architecture Decision Records {#appendix-b}\n\n")
    for path in ADRS:
        parts.append("\n\n")
        parts.append(_demote_first_heading(path.read_text(encoding="utf-8")))

    parts.append("\n\n# Appendix C. Reference index {#appendix-c}\n\n")
    reference_text = (ROOT / "references" / "REFERENCES.md").read_text(encoding="utf-8")
    reference_text = reference_text.replace("# References", "", 1)
    reference_text = rebase_relative_links(reference_text, "references/")
    parts.append(reference_text.strip() + "\n")

    out = ROOT / "MASTER_SPEC.md"
    out.write_text("".join(parts), encoding="utf-8")
    return out


def build_html(master: Path) -> Path:
    css = ROOT / "scripts" / "spec.css"
    out = ROOT / "MASTER_SPEC.html"
    cmd = [
        "pandoc",
        str(master),
        "--from=markdown+yaml_metadata_block",
        "--to=html5",
        "--standalone",
        "--embed-resources",
        f"--css={css}",
        "--metadata=pagetitle:sc-referee Scientific Audit Specification",
        "--output",
        str(out),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)
    return out


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _find_style(doc: Document, name: str):
    """Find a style by display name, including Pandoc heading styles with lookup quirks."""
    for style in doc.styles:
        if style.name == name:
            return style
    return None


def _ensure_style(doc: Document, name: str, base: str = "Normal"):
    styles = doc.styles
    existing = _find_style(doc, name)
    if existing is not None:
        return existing
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    base_style = _find_style(doc, base)
    if base_style is not None:
        style.base_style = base_style
    return style


def polish_docx(path: Path) -> None:
    doc = Document(path)
    doc.core_properties.title = ""  # LibreOffice otherwise duplicates the title on the cover.
    doc.core_properties.subject = "Architecture and product specification v0.5.0-draft"
    doc.core_properties.author = "sc-referee design project"
    doc.core_properties.keywords = "scientific audit, bioinformatics, workflow, Claude Code, specification"
    doc.core_properties.comments = "Generated review copy. Modular Markdown is the source of truth."

    navy = RGBColor(31, 67, 101)
    charcoal = RGBColor(31, 39, 48)
    muted = RGBColor(85, 96, 108)

    normal = _find_style(doc, "Normal")
    if normal is None:
        raise RuntimeError("DOCX lacks Normal style")
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(9.25)
    normal.font.color.rgb = charcoal
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.widow_control = True

    for name, size, before, after in [
        ("Title", 26, 0, 10),
        ("Subtitle", 14, 0, 14),
        ("Heading 1", 18, 8, 8),
        ("Heading 2", 14, 11, 5),
        ("Heading 3", 11.5, 8, 3),
        ("Heading 4", 10.5, 6, 2),
    ]:
        style = _find_style(doc, name)
        if style is None:
            continue
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = navy
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
    heading1 = _find_style(doc, "Heading 1")
    if heading1 is not None:
        heading1.paragraph_format.page_break_before = True
    subtitle = _find_style(doc, "Subtitle")
    if subtitle is not None:
        subtitle.font.color.rgb = muted

    code_style = _ensure_style(doc, "Source Code")
    code_style.font.name = "Liberation Mono"
    code_style.font.size = Pt(7.5)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.keep_together = False
    code_style.paragraph_format.widow_control = False

    for sty_name in ["Compact", "Block Text", "Body Text"]:
        style = _find_style(doc, sty_name)
        if style is not None:
            style.font.name = "Liberation Sans"
            style.font.size = Pt(9.0)

    # One compact, readable page geometry for the long review document.
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.68)
        section.right_margin = Inches(0.68)
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.25)
        section.different_first_page_header_footer = True
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "sc-referee Scientific Audit Specification  |  v0.5.0-draft"
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in hp.runs:
            run.font.name = "Liberation Sans"
            run.font.size = Pt(7.5)
            run.font.color.rgb = muted
        footer = section.footer
        fp = footer.paragraphs[0]
        _add_page_number(fp)
        for run in fp.runs:
            run.font.name = "Liberation Sans"
            run.font.size = Pt(8)
            run.font.color.rgb = muted

    # Pandoc tables: use the available width, repeat headers, and avoid oversized text.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            _set_repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                _set_cell_shading(cell, "E8EEF4")
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(1.5)
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.font.name = "Liberation Sans"
                        run.font.size = Pt(7.3 if len(table.columns) >= 5 else 8.0)
                        if r_idx == 0:
                            run.bold = True

    # Keep headings with following content; do not keep large code blocks/tables on one page.
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            para.paragraph_format.keep_with_next = True
        if para.style and para.style.name == "Source Code":
            para.paragraph_format.keep_together = False

    doc.save(path)


def build_docx(master: Path) -> Path:
    out = ROOT / DOCX_NAME
    cmd = [
        "pandoc",
        str(master),
        "--from=markdown+yaml_metadata_block",
        "--to=docx",
        "--metadata=pagetitle:sc-referee Scientific Audit Specification",
        "--output",
        str(out),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)
    polish_docx(out)
    return out


def update_document_index() -> None:
    candidates = (
        list((ROOT / "docs").glob("*.md"))
        + list((ROOT / "adrs").glob("*.md"))
        + list((ROOT / "templates").glob("*.md"))
        + list((ROOT / "examples").glob("*"))
        + list((ROOT / "references").glob("*.md"))
        + [
            ROOT / "README.md",
            ROOT / "CONTRIBUTING.md",
            ROOT / "REVIEW_CHECKLIST.md",
            ROOT / "ACCEPTANCE_CRITERIA.md",
            ROOT / "DECISIONS_v0.5.md",
            ROOT / "CHANGELOG.md",
        ]
    )
    documents = []
    for path in sorted(p for p in candidates if p.is_file()):
        text = path.read_text(encoding="utf-8", errors="replace")
        title = next((line[2:].strip() for line in text.splitlines() if line.startswith("# ")), path.stem)
        documents.append(
            {
                "path": str(path.relative_to(ROOT)),
                "title": title,
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                "word_count": len(re.findall(r"\b\w+[\w'-]*\b", text)),
            }
        )
    index = {
        "package": "sc-referee-specification",
        "version": VERSION,
        "schema_version": SCHEMA_VERSION,
        "date": DATE,
        "source_of_truth": "modular_markdown_adrs_and_machine_registers",
        "generated_review_copies": ["MASTER_SPEC.md", "MASTER_SPEC.html", DOCX_NAME],
        "documents": documents,
    }
    (ROOT / "machine" / "document-index.json").write_text(json.dumps(index, indent=2) + "\n", encoding="utf-8")


def build_manifest() -> Path:
    out = ROOT / "MANIFEST.sha256"
    ignored_names = {
        "MANIFEST.sha256",
        "VALIDATION.txt",
        "sc-referee-specification-v0.5.0-draft.zip",
        "sc-referee-specification-v0.5.0-draft.zip.sha256",
    }
    files = [
        p for p in ROOT.rglob("*")
        if (
            p.is_file()
            and p.name not in ignored_names
            and "__pycache__" not in p.parts
            and ".pytest_cache" not in p.parts
            and p.suffix != ".pyc"
        )
    ]
    lines = []
    for path in sorted(files):
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        lines.append(f"{digest}  {path.relative_to(ROOT)}")
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def main() -> None:
    master = build_master()
    html = build_html(master)
    docx = build_docx(master)
    update_document_index()
    manifest = build_manifest()
    print(
        "Built "
        + ", ".join(str(p.relative_to(ROOT)) for p in (master, html, docx, manifest))
    )


if __name__ == "__main__":
    main()
